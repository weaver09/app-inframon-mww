import os
import io
import uuid
import json

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    flash
)

from werkzeug.utils import secure_filename

from azure.identity import (
    DefaultAzureCredential,
    get_bearer_token_provider
)

from azure.storage.blob import BlobServiceClient
from openai import OpenAI
from mssql_python import connect
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook


# ============================================================
# Flask Application
# ============================================================

app = Flask(__name__)

app.secret_key = os.getenv(
    "FLASK_SECRET_KEY",
    "change-this-later"
)


# ============================================================
# Configuration
# ============================================================

SQL_CONNECTION_STRING = os.getenv(
    "AZURE_SQL_CONNECTIONSTRING"
)

STORAGE_ACCOUNT_NAME = os.getenv(
    "AZURE_STORAGE_ACCOUNT"
)

AZURE_OPENAI_ENDPOINT = os.getenv(
    "AZURE_OPENAI_ENDPOINT"
)

AZURE_OPENAI_DEPLOYMENT = os.getenv(
    "AZURE_OPENAI_DEPLOYMENT"
)

STORAGE_CONTAINER_NAME = "documents"

ALLOWED_EXTENSIONS = {
    "pdf",
    "docx",
    "xlsx"
}

MAX_FILE_SIZE_MB = 25

app.config["MAX_CONTENT_LENGTH"] = (
    MAX_FILE_SIZE_MB * 1024 * 1024
)


# ============================================================
# Database Connection
# ============================================================

def get_db_connection():

    if not SQL_CONNECTION_STRING:
        raise RuntimeError(
            "AZURE_SQL_CONNECTIONSTRING environment variable "
            "is not configured."
        )

    return connect(
        SQL_CONNECTION_STRING
    )


# ============================================================
# Azure Credential
# ============================================================

def get_azure_credential():

    return DefaultAzureCredential()


# ============================================================
# Azure Blob Storage
# ============================================================

def get_blob_service_client():

    if not STORAGE_ACCOUNT_NAME:
        raise RuntimeError(
            "AZURE_STORAGE_ACCOUNT environment variable "
            "is not configured."
        )

    account_url = (
        f"https://{STORAGE_ACCOUNT_NAME}.blob.core.windows.net"
    )

    credential = get_azure_credential()

    return BlobServiceClient(
        account_url=account_url,
        credential=credential
    )


# ============================================================
# Azure OpenAI
# ============================================================

def get_openai_client():

    if not AZURE_OPENAI_ENDPOINT:
        raise RuntimeError(
            "AZURE_OPENAI_ENDPOINT environment variable "
            "is not configured."
        )

    if not AZURE_OPENAI_DEPLOYMENT:
        raise RuntimeError(
            "AZURE_OPENAI_DEPLOYMENT environment variable "
            "is not configured."
        )

    credential = DefaultAzureCredential()

    token_provider = get_bearer_token_provider(
        credential,
        "https://cognitiveservices.azure.com/.default"
    )

    base_url = (
        AZURE_OPENAI_ENDPOINT.rstrip("/")
        + "/openai/v1/"
    )

    return OpenAI(
        base_url=base_url,
        api_key=token_provider
    )


# ============================================================
# File Validation
# ============================================================

def allowed_file(filename):

    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower()
        in ALLOWED_EXTENSIONS
    )


# ============================================================
# PDF Extraction
# ============================================================

def extract_pdf(file_data):

    reader = PdfReader(
        io.BytesIO(file_data)
    )

    parts = []

    for page_number, page in enumerate(
        reader.pages,
        start=1
    ):

        text = page.extract_text() or ""

        if text.strip():
            parts.append(
                f"\n--- Page {page_number} ---\n{text}"
            )

    return "\n".join(parts)


# ============================================================
# DOCX Extraction
# ============================================================

def extract_docx(file_data):

    document = Document(
        io.BytesIO(file_data)
    )

    parts = []

    # Paragraphs
    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # Tables
    for table_index, table in enumerate(
        document.tables,
        start=1
    ):

        parts.append(
            f"\n--- Table {table_index} ---"
        )

        for row in table.rows:

            values = [
                cell.text.strip()
                for cell in row.cells
            ]

            parts.append(
                " | ".join(values)
            )

    return "\n".join(parts)


# ============================================================
# XLSX Extraction
# ============================================================

def extract_xlsx(file_data):

    workbook = load_workbook(
        io.BytesIO(file_data),
        read_only=True,
        data_only=True
    )

    parts = []

    for worksheet in workbook.worksheets:

        parts.append(
            f"\n--- Worksheet: {worksheet.title} ---"
        )

        for row in worksheet.iter_rows(
            values_only=True
        ):

            values = []

            for value in row:

                if value is None:
                    values.append("")
                else:
                    values.append(
                        str(value)
                    )

            if any(
                value.strip()
                for value in values
            ):
                parts.append(
                    " | ".join(values)
                )

    workbook.close()

    return "\n".join(parts)


# ============================================================
# Extract Content
# ============================================================

def extract_content(
    filename,
    file_data
):

    extension = (
        filename
        .lower()
        .rsplit(".", 1)[-1]
    )

    if extension == "pdf":
        return extract_pdf(
            file_data
        )

    if extension == "docx":
        return extract_docx(
            file_data
        )

    if extension == "xlsx":
        return extract_xlsx(
            file_data
        )

    raise ValueError(
        f"Unsupported file type: {extension}"
    )


# ============================================================
# AI Document Analysis
# ============================================================

def analyze_document_with_ai(
    extracted_text
):

    if not extracted_text:
        raise RuntimeError(
            "No extracted text was available "
            "for AI analysis."
        )

    client = get_openai_client()

    # Keep initial calls reasonably sized
    max_characters = 60000

    text_to_analyze = (
        extracted_text[:max_characters]
    )

    system_message = """
You are an AI document analysis system.

Analyze documents accurately and conservatively.

Do not invent information.

Return only valid JSON.

If information is not present in the document,
use an empty array or null value where appropriate.
"""

    user_message = f"""
Analyze the following document.

Return ONLY valid JSON using exactly this structure:

{{
    "summary": "Concise but useful document summary",

    "key_topics": [
        "topic"
    ],

    "entities": [
        {{
            "name": "entity name",
            "type": "person|organization|product|location|system|other"
        }}
    ],

    "tags": [
        "tag"
    ],

    "risks": [
        {{
            "risk": "risk description",
            "severity": "Low|Medium|High"
        }}
    ],

    "action_items": [
        {{
            "action": "action description",
            "owner": null,
            "due_date": null
        }}
    ],

    "important_dates": [
        {{
            "date": "YYYY-MM-DD or original date text",
            "description": "why this date matters"
        }}
    ]
}}

Rules:

1. Do not invent facts.
2. Summary should be factual and useful.
3. Key topics should identify major themes.
4. Entities should identify important people,
   organizations, systems, products, and locations.
5. Tags should be short and useful for searching.
6. Risks must be supported by the document.
7. If no risks exist, return an empty array.
8. Action items must come from explicit or strongly
   implied tasks in the document.
9. Preserve owners when explicitly stated.
10. Preserve due dates when explicitly stated.
11. If no action items exist, return an empty array.
12. Important dates should only include meaningful dates.
13. If no important dates exist, return an empty array.

DOCUMENT:

{text_to_analyze}
"""

    response = client.chat.completions.create(

        model=AZURE_OPENAI_DEPLOYMENT,

        messages=[
            {
                "role": "system",
                "content": system_message
            },
            {
                "role": "user",
                "content": user_message
            }
        ],

        temperature=0.1,

        response_format={
            "type": "json_object"
        }
    )

    content = (
        response
        .choices[0]
        .message
        .content
    )

    if not content:
        raise RuntimeError(
            "Azure OpenAI returned "
            "an empty response."
        )

    try:

        analysis = json.loads(
            content
        )

    except json.JSONDecodeError as exc:

        raise RuntimeError(
            "Azure OpenAI returned "
            "invalid JSON."
        ) from exc

    # Make sure every expected property exists
    analysis.setdefault(
        "summary",
        ""
    )

    analysis.setdefault(
        "key_topics",
        []
    )

    analysis.setdefault(
        "entities",
        []
    )

    analysis.setdefault(
        "tags",
        []
    )

    analysis.setdefault(
        "risks",
        []
    )

    analysis.setdefault(
        "action_items",
        []
    )

    analysis.setdefault(
        "important_dates",
        []
    )

    return analysis


# ============================================================
# Home Page
# ============================================================

@app.route("/")
def index():

    documents = []
    error_message = None

    conn = None
    cursor = None

    try:

        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT
                d.DocumentID,
                d.FileName,
                d.BlobName,
                d.FileType,
                d.FileSizeBytes,
                d.UploadDate,
                d.Status,
                LEN(d.ExtractedText)
                    AS ExtractedCharacters,
                a.AnalysisID,
                a.Summary
            FROM dbo.Documents d
            LEFT JOIN dbo.DocumentAnalysis a
                ON d.DocumentID = a.DocumentID
            ORDER BY
                d.UploadDate DESC;
        """)

        rows = cursor.fetchall()

        for row in rows:

            documents.append({
                "DocumentID": row[0],
                "FileName": row[1],
                "BlobName": row[2],
                "FileType": row[3],
                "FileSizeBytes": row[4],
                "UploadDate": row[5],
                "Status": row[6],
                "ExtractedCharacters": row[7],
                "AnalysisID": row[8],
                "Summary": row[9]
            })

    except Exception as exc:

        error_message = str(exc)

    finally:

        try:
            if cursor:
                cursor.close()
        except Exception:
            pass

        try:
            if conn:
                conn.close()
        except Exception:
            pass

    return render_template(
        "index.html",
        documents=documents,
        error_message=error_message
    )


# ============================================================
# Upload Document
# ============================================================

@app.route(
    "/upload",
    methods=["POST"]
)
def upload_document():

    if "file" not in request.files:

        flash(
            "No file was selected.",
            "error"
        )

        return redirect(
            url_for("index")
        )

    file = request.files["file"]

    if file.filename == "":

        flash(
            "No file was selected.",
            "error"
        )

        return redirect(
            url_for("index")
        )

    if not allowed_file(
        file.filename
    ):

        flash(
            "Unsupported file type. "
            "Only PDF, DOCX, and XLSX "
            "files are allowed.",
            "error"
        )

        return redirect(
            url_for("index")
        )

    original_filename = (
        secure_filename(
            file.filename
        )
    )

    extension = (
        original_filename
        .rsplit(".", 1)[1]
        .lower()
    )

    unique_blob_name = (
        f"{uuid.uuid4()}-"
        f"{original_filename}"
    )

    document_id = None
    conn = None
    cursor = None

    try:

        # ====================================================
        # Read uploaded file
        # ====================================================

        file_data = file.read()

        file_size = len(
            file_data
        )

        # ====================================================
        # Upload to Blob Storage
        # ====================================================

        blob_service_client = (
            get_blob_service_client()
        )

        blob_client = (
            blob_service_client
            .get_blob_client(
                container=STORAGE_CONTAINER_NAME,
                blob=unique_blob_name
            )
        )

        blob_client.upload_blob(
            file_data,
            overwrite=False
        )

        # ====================================================
        # Insert metadata into SQL
        # ====================================================

        conn = get_db_connection()

        cursor = conn.cursor()

        cursor.execute("""
            INSERT INTO dbo.Documents
            (
                FileName,
                BlobName,
                FileType,
                FileSizeBytes,
                Status
            )
            OUTPUT
                INSERTED.DocumentID
            VALUES
            (
                ?,
                ?,
                ?,
                ?,
                ?
            );
        """,
        (
            original_filename,
            unique_blob_name,
            extension,
            file_size,
            "Processing"
        ))

        row = cursor.fetchone()

        document_id = row[0]

        conn.commit()

        # ====================================================
        # Extract document content
        # ====================================================

        extracted_text = (
            extract_content(
                original_filename,
                file_data
            )
        )

        # ====================================================
        # Save extracted text
        # ====================================================

        cursor.execute("""
            UPDATE dbo.Documents
            SET
                ExtractedText = ?,
                Status = ?
            WHERE
                DocumentID = ?;
        """,
        (
            extracted_text,
            "Extracted",
            document_id
        ))

        conn.commit()

        # ====================================================
        # Azure OpenAI analysis
        # ====================================================

        analysis = (
            analyze_document_with_ai(
                extracted_text
            )
        )

        # ====================================================
        # Store AI results
        # ====================================================

        cursor.execute("""
            INSERT INTO dbo.DocumentAnalysis
            (
                DocumentID,
                Summary,
                KeyTopics,
                EntitiesJson,
                TagsJson,
                RisksJson,
                ActionItemsJson,
                ImportantDatesJson
            )
            VALUES
            (
                ?,
                ?,
                ?,
                ?,
                ?,
                ?,
                ?,
                ?
            );
        """,
        (
            document_id,

            analysis.get(
                "summary"
            ),

            json.dumps(
                analysis.get(
                    "key_topics",
                    []
                )
            ),

            json.dumps(
                analysis.get(
                    "entities",
                    []
                )
            ),

            json.dumps(
                analysis.get(
                    "tags",
                    []
                )
            ),

            json.dumps(
                analysis.get(
                    "risks",
                    []
                )
            ),

            json.dumps(
                analysis.get(
                    "action_items",
                    []
                )
            ),

            json.dumps(
                analysis.get(
                    "important_dates",
                    []
                )
            )
        ))

        # ====================================================
        # Mark document analyzed
        # ====================================================

        cursor.execute("""
            UPDATE dbo.Documents
            SET
                Status = ?
            WHERE
                DocumentID = ?;
        """,
        (
            "Analyzed",
            document_id
        ))

        conn.commit()

        flash(
            f"{original_filename} "
            "uploaded, extracted, and "
            "analyzed successfully.",
            "success"
        )

    except Exception as exc:

        # ====================================================
        # Mark failed document
        # ====================================================

        try:

            if document_id:

                if conn is None:
                    conn = (
                        get_db_connection()
                    )

                if cursor is None:
                    cursor = (
                        conn.cursor()
                    )

                cursor.execute("""
                    UPDATE dbo.Documents
                    SET
                        Status = ?
                    WHERE
                        DocumentID = ?;
                """,
                (
                    "Processing Failed",
                    document_id
                ))

                conn.commit()

        except Exception:
            pass

        flash(
            f"Upload failed: "
            f"{str(exc)}",
            "error"
        )

    finally:

        try:
            if cursor:
                cursor.close()
        except Exception:
            pass

        try:
            if conn:
                conn.close()
        except Exception:
            pass

    return redirect(
        url_for("index")
    )


# ============================================================
# Health Check
# ============================================================

@app.route("/health")
def health():

    results = {
        "database": "unknown",
        "storage": "unknown",
        "openai": "unknown"
    }

    status_code = 200

    # ========================================================
    # SQL Test
    # ========================================================

    try:

        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute(
            "SELECT 1;"
        )

        cursor.fetchone()

        cursor.close()
        conn.close()

        results[
            "database"
        ] = "connected"

    except Exception as exc:

        results[
            "database"
        ] = (
            f"error: {str(exc)}"
        )

        status_code = 500

    # ========================================================
    # Blob Storage Test
    # ========================================================

    try:

        blob_service_client = (
            get_blob_service_client()
        )

        container_client = (
            blob_service_client
            .get_container_client(
                STORAGE_CONTAINER_NAME
            )
        )

        container_client.get_container_properties()

        results[
            "storage"
        ] = "connected"

    except Exception as exc:

        results[
            "storage"
        ] = (
            f"error: {str(exc)}"
        )

        status_code = 500

    # ========================================================
    # Azure OpenAI Configuration Test
    # ========================================================

    try:

        if not AZURE_OPENAI_ENDPOINT:

            raise RuntimeError(
                "AZURE_OPENAI_ENDPOINT "
                "is not configured."
            )

        if not AZURE_OPENAI_DEPLOYMENT:

            raise RuntimeError(
                "AZURE_OPENAI_DEPLOYMENT "
                "is not configured."
            )

        get_openai_client()

        results[
            "openai"
        ] = "configured"

    except Exception as exc:

        results[
            "openai"
        ] = (
            f"error: {str(exc)}"
        )

        status_code = 500

    # ========================================================
    # Overall Status
    # ========================================================

    results["status"] = (
        "healthy"
        if status_code == 200
        else "unhealthy"
    )

    return (
        results,
        status_code
    )


# ============================================================
# Application Entry Point
# ============================================================

if __name__ == "__main__":

    app.run(
        host="0.0.0.0",
        port=8000
    )